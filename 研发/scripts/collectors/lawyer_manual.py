from __future__ import annotations

import argparse
import asyncio
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document

ROOT_DIR = Path(__file__).resolve().parents[2]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from scripts.collectors.common import (  # noqa: E402
    RAW_SOURCES_DIR,
    ensure_dir,
    safe_slug,
    stable_doc_id,
    write_collected_document,
)

SOURCE_DIR = RAW_SOURCES_DIR / "lawyer_manual"
COLLECTED_ROLE_ID = "lawyer_01"
DATE_PATTERN = re.compile(r"(20\d{6}|\d{8})$")


class LawyerManualCollector:
    def __init__(self, *, max_docs: int | None = None) -> None:
        self.max_docs = max_docs
        self.role_id = COLLECTED_ROLE_ID
        self.input_roots = [
            ("statutes", "official", "statute"),
            ("judicial_interpretations", "official", "judicial_interpretation"),
            ("cases", "case_law", "case"),
        ]
        ensure_dir(SOURCE_DIR)

    async def collect(self) -> list[dict[str, Any]]:
        docs: list[dict[str, Any]] = []
        for folder_name, source_tier, content_type in self.input_roots:
            folder = SOURCE_DIR / folder_name
            if not folder.exists():
                continue
            for file_path in sorted(folder.iterdir()):
                if self.max_docs is not None and len(docs) >= self.max_docs:
                    return docs
                if not file_path.is_file():
                    continue
                suffix = file_path.suffix.lower()
                if suffix == ".docx":
                    doc = await asyncio.to_thread(
                        self._parse_docx,
                        file_path,
                        folder_name,
                        source_tier,
                        content_type,
                    )
                    if doc:
                        write_collected_document(self.role_id, doc)
                        docs.append(doc)
                elif suffix == ".doc":
                    continue
        return docs

    def _parse_docx(
        self,
        file_path: Path,
        folder_name: str,
        source_tier: str,
        content_type: str,
    ) -> dict[str, Any] | None:
        document = Document(str(file_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_lines: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))

        content_parts = []
        if paragraphs:
            content_parts.append("\n".join(paragraphs))
        if table_lines:
            content_parts.append("\n".join(table_lines))
        content = "\n\n".join(part for part in content_parts if part).strip()
        if len(content) < 80:
            return None

        title, published_at = self._parse_name_metadata(file_path.stem)
        relative_path = file_path.relative_to(SOURCE_DIR).as_posix()
        source_url = f"local://lawyer_manual/{relative_path}"
        doc_id = stable_doc_id("lawyer", source_url)

        return {
            "doc_id": doc_id,
            "title": title,
            "content": content,
            "source_name": "lawyer_manual",
            "source_url": source_url,
            "source_domain": "local",
            "published_at": published_at,
            "role_id": self.role_id,
            "source_tier": source_tier,
            "tags": [self.role_id, folder_name, content_type],
            "metadata": {
                "collector": "lawyer_manual",
                "content_type": content_type,
                "source_bucket": folder_name,
                "original_file_name": file_path.name,
                "original_extension": file_path.suffix.lower(),
                "safe_slug": safe_slug(title),
            },
        }

    def _parse_name_metadata(self, stem: str) -> tuple[str, str | None]:
        title = stem
        published_at: str | None = None
        if "_" not in stem:
            return title, published_at

        prefix, maybe_date = stem.rsplit("_", 1)
        if DATE_PATTERN.fullmatch(maybe_date):
            title = prefix
            if len(maybe_date) == 8:
                published_at = f"{maybe_date[:4]}-{maybe_date[4:6]}-{maybe_date[6:8]}"
        return title, published_at


async def _main_async(args: argparse.Namespace) -> None:
    collector = LawyerManualCollector(max_docs=args.max_docs)
    docs = await collector.collect()
    print(f"lawyer_01 collected {len(docs)} documents")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--max-docs", type=int, default=None)
    args = parser.parse_args()
    asyncio.run(_main_async(args))


if __name__ == "__main__":
    main()
